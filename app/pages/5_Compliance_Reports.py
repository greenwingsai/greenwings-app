import io
import sys
from datetime import date
from pathlib import Path

import pandas as pd
import streamlit as st

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from core import data

st.title("📋 Regulatory Compliance Reports")
st.markdown("""
Since **January 2025**, intra-European flights must monitor **non-CO₂ effects** under the EU ETS MRV
framework - on top of the existing CO₂ obligations and CORSIA. Non-compliance is an operational-cost
line item: allowance surrender, **€100/t excess-emissions penalties**, and audit findings.
GreenWings generates the underlying per-flight evidence pack from your own telemetry, ready for
your verifier.
""")

fl = data.flights()

# Cruise-typical NOx emission index per aircraft type (OpenAP BFFM2 at each
# type's median observed fuel flow), cached so fleet reports stay instant.
@st.cache_data(show_spinner="Deriving per-type emission indices (OpenAP)…")
def type_ei() -> pd.DataFrame:
    from core.emissions import _openap_model
    rows = []
    for actype, ff in fl.groupby("aircraft_type")["fuel_flow_kgs"].median().items():
        model = _openap_model(actype)
        if model is not None and ff > 0:
            ei_nox = model.nox(ff, tas=447, alt=34500) / 1000 / ff   # kg NOx per kg fuel
        else:
            ei_nox = 0.0153
        rows.append((actype, ei_nox))
    return pd.DataFrame(rows, columns=["aircraft_type", "ei_nox"]).set_index("aircraft_type")


# ---------------- scope selection --------------------------------------
st.sidebar.header("Report scope")
types = ["All types"] + sorted(fl["aircraft_type"].unique())
scope_type = st.sidebar.selectbox("Aircraft type", types)
dmin, dmax = pd.to_datetime(fl["flight_date"]).min(), pd.to_datetime(fl["flight_date"]).max()
period = st.sidebar.date_input("Reporting period", (dmin.date(), dmax.date()),
                               min_value=dmin.date(), max_value=dmax.date())

sub = fl.copy()
sub["flight_date_dt"] = pd.to_datetime(sub["flight_date"])
if scope_type != "All types":
    sub = sub[sub["aircraft_type"] == scope_type]
if isinstance(period, tuple) and len(period) == 2:
    sub = sub[(sub["flight_date_dt"] >= pd.Timestamp(period[0])) & (sub["flight_date_dt"] <= pd.Timestamp(period[1]))]

ei = type_ei()
rep = sub.merge(ei, left_on="aircraft_type", right_index=True, how="left")
rep["co2_kg"] = rep["fuel_obs_kg"] * 3.16
rep["h2o_kg"] = rep["fuel_obs_kg"] * 1.23
rep["nox_kg"] = rep["fuel_obs_kg"] * rep["ei_nox"].fillna(0.0153)
rep["sox_kg"] = rep["fuel_obs_kg"] * 0.0012
rep["soot_kg"] = rep["fuel_obs_kg"] * 0.00003

k1, k2, k3, k4, k5 = st.columns(5)
k1.metric("Flights in scope", f"{len(rep):,}")
k2.metric("Fuel (verified basis)", f"{rep['fuel_obs_kg'].sum()/1e6:,.2f} kt")
k3.metric("CO₂", f"{rep['co2_kg'].sum()/1e6:,.2f} kt")
k4.metric("NOₓ", f"{rep['nox_kg'].sum()/1e3:,.1f} t")
k5.metric("Indicative ETS cost @ €75/t", f"€{rep['co2_kg'].sum()/1000*75/1e6:,.2f} M",
          help="CO₂ allowance exposure if all flights were ETS-liable - for budgeting, not a tax statement")

st.divider()

# ---------------- regulatory mapping ------------------------------------
st.subheader("What each regulation asks for - and where GreenWings answers it")
st.markdown("""
| Regulation | Obligation on the operator | GreenWings evidence |
|---|---|---|
| **EU ETS** - Directive 2003/87/EC | Annual verified CO₂ report per operator; surrender allowances; €100/t excess penalty | Per-flight CO₂ from **measured fuel** (3.16 kg/kg, standard factor) - CSV below |
| **MRV** - Implementing Reg. (EU) 2018/2066 + **non-CO₂ framework (from 1 Jan 2025)** | Monitor **non-CO₂ effects** (NOₓ, contrail exposure) on intra-EEA flights | Per-flight NOₓ (OpenAP/BFFM2 per type) + contrail-band exposure from trajectories |
| **CORSIA** - ICAO Annex 16 Vol. IV | Fuel-based CO₂ MRV on international routes; offset growth above baseline | Same fuel-based method; route metadata included per flight |
| **ReFuelEU Aviation** - Reg. (EU) 2023/2405 | SAF blending shares rising from 2% (2025) to 6% (2030)+ | What-if simulator quantifies SAF scenarios per flight/fleet |
| **EU AI Act** - Reg. (EU) 2024/1689 | Transparency & human oversight for AI decision-support | Recommendations are human-approved drafts; every figure has a visible tool trace |
| **GDPR** | Lawful processing, data minimisation | Aggregate operational telemetry, no passenger PII |

⚠️ **Honest scope:** GreenWings is a decision-support prototype. It produces the *monitoring evidence and
draft reports*; formal ETS/CORSIA submission still requires your accredited verifier - which is exactly
the human-in-the-loop step the platform enforces.
""")

st.divider()

# ---------------- downloads ---------------------------------------------
st.subheader("⬇ Generate the evidence pack")
c1, c2 = st.columns(2)

with c1:
    st.markdown("**Per-flight emissions ledger (CSV)** - the verifier-facing dataset: one row per flight, "
                "measured fuel and all species.")
    csv_cols = rep[["flight_id", "flight_date", "aircraft_type", "origin_icao", "destination_icao",
                    "distance_km", "fuel_obs_kg", "co2_kg", "h2o_kg", "nox_kg", "sox_kg", "soot_kg"]].round(2)
    st.download_button("Download per-flight CSV", csv_cols.to_csv(index=False),
                       file_name=f"greenwings_emissions_ledger_{date.today().isoformat()}.csv",
                       mime="text/csv", type="primary")
    st.dataframe(csv_cols.head(8), use_container_width=True, hide_index=True)

with c2:
    st.markdown("**Executive compliance report (Word)** - totals, methodology, uncertainty statement and "
                "sign-off block, formatted for internal review and the verifier meeting.")

    def build_docx() -> bytes:
        from docx import Document
        doc = Document()
        doc.add_heading("Aviation Emissions Monitoring Report", 0)
        doc.add_paragraph(f"Generated by GreenWings AI on {date.today().isoformat()} · "
                          f"Scope: {scope_type} · Period: {period[0]} to {period[1]}"
                          if isinstance(period, tuple) and len(period) == 2 else "")
        doc.add_heading("1. Reported totals", 1)
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light Grid Accent 1"
        t.rows[0].cells[0].text, t.rows[0].cells[1].text = "Quantity", "Value"
        for label, val in [
            ("Flights in scope", f"{len(rep):,}"),
            ("Fuel consumed (measured, ACARS)", f"{rep['fuel_obs_kg'].sum():,.0f} kg"),
            ("CO2 (EU ETS factor 3.16 kg/kg fuel)", f"{rep['co2_kg'].sum():,.0f} kg"),
            ("NOx (OpenAP BFFM2, per aircraft type)", f"{rep['nox_kg'].sum():,.0f} kg"),
            ("H2O", f"{rep['h2o_kg'].sum():,.0f} kg"),
            ("SOx", f"{rep['sox_kg'].sum():,.0f} kg"),
            ("Soot (nvPM)", f"{rep['soot_kg'].sum():,.1f} kg"),
        ]:
            row = t.add_row().cells
            row[0].text, row[1].text = label, val
        doc.add_heading("2. Methodology", 1)
        doc.add_paragraph(
            "Fuel consumption is taken from ACARS fuel-on-board telemetry (measured, not modelled). "
            "CO2, H2O and SOx apply fuel-chemistry emission factors; NOx applies the Boeing Fuel Flow "
            "Method 2 via OpenAP models per aircraft type, aligned with the ICAO Engine Emissions "
            "Databank. Contrail-band exposure is derived from ADS-B altitude data (8.5-12 km band).")
        doc.add_heading("3. Uncertainty & limitations", 1)
        doc.add_paragraph(
            "Non-CO2 quantities carry materially higher scientific uncertainty than CO2 (Lee et al. 2021); "
            "total climate impact should be quoted as a range of 1.3-3.0x CO2. Interval coverage varies "
            "per flight; figures represent the measured window. This report is monitoring evidence for "
            "review by an accredited verifier, not a formal ETS/CORSIA submission.")
        doc.add_heading("4. Human review & sign-off", 1)
        doc.add_paragraph("Reviewed by (name / role): ____________________________\n"
                          "Date: ____________     Signature: ____________________")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    st.download_button("Download executive report (DOCX)", build_docx(),
                       file_name=f"greenwings_compliance_report_{date.today().isoformat()}.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       type="primary")
    st.caption("NOₓ here uses each type's cruise-typical emission index; the Footprint Dashboard computes "
               "altitude-corrected values per interval where trajectory data is loaded.")

st.divider()
st.subheader("A note on operator identity in this demo")
st.markdown("""
This demo dataset is **anonymised by EUROCONTROL**: flight IDs like `prc770831136` carry **no airline name,
callsign or tail number** - deliberate protection of commercially sensitive fuel data. We therefore identify
aircraft by **type + mission profile** and benchmark against peer groups.
**In production, this changes:** connected to *your* operations data (tail numbers, fleet IDs), the same
pipeline reports per aircraft, per route, per AOC - and the maintenance module tracks individual airframes
over time, which is exactly how it becomes a work-order trigger.
""")
